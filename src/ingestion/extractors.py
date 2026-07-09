"""Extração de texto de documentos jurídicos (PDF, DOCX, TXT/MD).

Além da extração bruta, aplica a limpeza descrita na metodologia:
- remoção de cabeçalhos/rodapés repetidos entre páginas;
- remoção de linhas de numeração de página;
- normalização de espaços em branco.

OCR (opcional): quando `ocr_enabled=True` e `pytesseract` está instalado,
páginas de PDF sem texto selecionável são rasterizadas e submetidas a OCR
em português — evolução sobre o design original, que não tratava PDFs
escaneados.
"""

from __future__ import annotations

import io
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path


@dataclass
class PageText:
    """Texto extraído de uma página (1-indexada) de um documento."""

    page: int
    text: str


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

_PAGENUM_RE = re.compile(r"^\s*(p[áa]g(ina)?\.?\s*)?\d+(\s*(de|/)\s*\d+)?\s*$", re.IGNORECASE)


class UnsupportedFormatError(ValueError):
    """Formato de arquivo não suportado pela ingestão."""


def extract_document(filename: str, data: bytes, *, ocr_enabled: bool = False) -> list[PageText]:
    """Extrai o texto de `data` de acordo com a extensão de `filename`."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        pages = _extract_pdf(data, ocr_enabled=ocr_enabled)
    elif ext == ".docx":
        pages = _extract_docx(data)
    elif ext in {".txt", ".md"}:
        pages = [PageText(page=1, text=data.decode("utf-8", errors="replace"))]
    else:
        raise UnsupportedFormatError(
            f"Extensão '{ext}' não suportada. Formatos aceitos: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    return remove_boilerplate([PageText(p.page, normalize_text(p.text)) for p in pages])


def _extract_pdf(data: bytes, *, ocr_enabled: bool) -> list[PageText]:
    import fitz  # PyMuPDF

    pages: list[PageText] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if not text and ocr_enabled:
                text = _ocr_page(page)
            pages.append(PageText(page=i, text=text))
    return pages


def _ocr_page(page) -> str:
    """OCR de uma página rasterizada; retorna '' se o OCR não estiver disponível."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""
    pix = page.get_pixmap(dpi=200)
    image = Image.open(io.BytesIO(pix.tobytes("png")))
    try:
        return pytesseract.image_to_string(image, lang="por")
    except Exception:
        # Binário do tesseract ausente ou sem o pacote de idioma 'por'.
        return ""


def _extract_docx(data: bytes) -> list[PageText]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    # DOCX não tem paginação fixa: todo o conteúdo é tratado como página 1.
    return [PageText(page=1, text="\n".join(parts))]


def normalize_text(text: str) -> str:
    """Normaliza espaços: remove espaços à direita e colapsa linhas em branco."""
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def remove_boilerplate(
    pages: list[PageText], *, min_pages: int = 3, repeat_ratio: float = 0.6, max_line_len: int = 80
) -> list[PageText]:
    """Remove cabeçalhos/rodapés (linhas curtas repetidas na maioria das páginas)
    e linhas que são apenas numeração de página."""
    repeated: set[str] = set()
    if len(pages) >= min_pages:
        counts: Counter[str] = Counter()
        for p in pages:
            unique_lines = {
                line.strip()
                for line in p.text.splitlines()
                if 0 < len(line.strip()) <= max_line_len
            }
            counts.update(unique_lines)
        repeated = {line for line, c in counts.items() if c / len(pages) >= repeat_ratio}

    cleaned: list[PageText] = []
    for p in pages:
        lines = [
            line
            for line in p.text.splitlines()
            if line.strip() not in repeated and not _PAGENUM_RE.match(line)
        ]
        cleaned.append(PageText(page=p.page, text="\n".join(lines).strip()))
    return cleaned
