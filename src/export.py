"""Exportação de respostas para Word (.docx).

Funcionalidade listada como "trabalho futuro" no design original e entregue
aqui: converte o Markdown leve das respostas (títulos, listas, negrito) em um
documento Word com cabeçalho e aviso legal.
"""

from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt

DISCLAIMER = (
    "Documento gerado pelo Themis.IA como apoio à atividade jurídica. "
    "As teses e redações não substituem a análise humana rigorosa de um(a) "
    "advogado(a) responsável."
)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_paragraph(doc: Document, line: str) -> None:
    """Adiciona uma linha de Markdown leve como parágrafo do Word."""
    stripped = line.strip()
    if not stripped:
        return

    heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if heading:
        level = min(len(heading.group(1)) + 1, 4)
        doc.add_heading(_BOLD_RE.sub(r"\1", heading.group(2)), level=level)
        return

    style = None
    if re.match(r"^[-*•]\s+", stripped):
        stripped = re.sub(r"^[-*•]\s+", "", stripped)
        style = "List Bullet"
    elif re.match(r"^\d+[.)]\s+", stripped):
        stripped = re.sub(r"^\d+[.)]\s+", "", stripped)
        style = "List Number"

    paragraph = doc.add_paragraph(style=style)
    # Alterna trechos normais e em negrito delimitados por **...**
    pieces = _BOLD_RE.split(stripped)
    for i, piece in enumerate(pieces):
        if not piece:
            continue
        run = paragraph.add_run(piece)
        run.bold = i % 2 == 1


def answer_to_docx_bytes(title: str, content: str, persona_label: str) -> bytes:
    """Gera um .docx (bytes) a partir de uma resposta do assistente."""
    doc = Document()

    doc.add_heading("Themis.IA — Minuta de Apoio", level=0)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Assunto: {title}\nPersona: {persona_label}\n"
        f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.italic = True

    for line in content.splitlines():
        _add_markdown_paragraph(doc, line)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(DISCLAIMER)
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
